#!/usr/bin/env python
'''
various utilities for formatting a docx
'''

import re
import sys
import lxml
import types
import string
import random
import requests
import importlib
import traceback

from copy import deepcopy
from pathlib import Path
from PIL import Image
from lxml import etree

from docx import Document, section, document, table
from docx.oxml import OxmlElement, parse_xml, ns
from docx.oxml.ns import qn, nsdecls

from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt, Cm, Inches, RGBColor

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.text.paragraph import Paragraph

import latex2mathml.converter

from googleapiclient import errors
from googleapiclient.http import MediaIoBaseDownload

from ggle.google_services import GoogleServices
from helper.config_service import ConfigService
from helper.logger import *

if sys.platform in ['win32', 'darwin']:
	import win32com.client as client


# --------------------------------------------------------------------------------------------------------------------------------------------
# document and section related

''' process a list of section_data and generate docx code
'''
def section_list_to_docx(docx, section_list, nesting_level=0):
	first_section = True
	for section_data in section_list:
		section_meta = section_data['section-meta']
		section_prop = section_data['section-prop']

		if section_prop['label'] != '':
			info(f"writing : {section_prop['label'].strip()} {section_prop['heading'].strip()}", nesting_level=nesting_level)
		else:
			info(f"writing : {section_prop['heading'].strip()}", nesting_level=nesting_level)

		module = importlib.import_module("doc.docx_api")
		func = getattr(module, f"process_{section_prop['content-type']}")
		func(docx=docx, section_data=section_data, nesting_level=nesting_level+1)


''' add or update a document section
'''
def add_or_update_document_section(docx, page_spec, margin_spec, orientation, different_firstpage, section_break, page_break, first_section, different_odd_even_pages, link_to_previous=False, nesting_level=0):
	#  if it is a section break, we isnert a new section
	if section_break:
		new_section = True
		docx_section = docx.add_section(WD_SECTION.NEW_PAGE)

	else:
		# we are continuing the last section
		if first_section:
			new_section = True
		else:
			new_section = False

		docx_section = docx.sections[-1]

		#  we may have a page break
		if page_break:
			docx.add_page_break()


	docx_section.first_page_header.is_linked_to_previous = link_to_previous
	docx_section.header.is_linked_to_previous = link_to_previous
	docx_section.even_page_header.is_linked_to_previous = link_to_previous

	docx_section.first_page_footer.is_linked_to_previous = link_to_previous
	docx_section.footer.is_linked_to_previous = link_to_previous
	docx_section.even_page_footer.is_linked_to_previous = link_to_previous

	if orientation == 'landscape':
		docx_section.orient = WD_ORIENT.LANDSCAPE
		docx_section.page_width = Inches(page_spec['height'])
		docx_section.page_height = Inches(page_spec['width'])
	else:
		docx_section.orient = WD_ORIENT.PORTRAIT
		docx_section.page_width = Inches(page_spec['width'])
		docx_section.page_height = Inches(page_spec['height'])

	docx_section.left_margin = Inches(margin_spec['left'])
	docx_section.right_margin = Inches(margin_spec['right'])
	docx_section.top_margin = Inches(margin_spec['top'])
	docx_section.bottom_margin = Inches(margin_spec['bottom'])

	docx_section.gutter = Inches(margin_spec['gutter'])

	docx_section.header_distance = Inches(margin_spec['distance']['header'])
	docx_section.footer_distance = Inches(margin_spec['distance']['footer'])

	docx_section.different_first_page_header_footer = different_firstpage
	docx.settings.odd_and_even_pages_header_footer = different_odd_even_pages

	# get the actual width
	# actual_width = docx_section.page_width.inches - docx_section.left_margin.inches - docx_section.right_margin.inches - docx_section.gutter.inches

	return docx_section, new_section



# -----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# various utility functions

''' format container (paragraph or cell)
	border, bgcolor, padding, valign, halign
'''
def format_container(container, attributes, custom_style_name=None, it_is_a_table_cell=True, nesting_level=0):
	custom_style = ConfigService()._style_specs.get(custom_style_name, None)
	if custom_style:
		if 'padding' in custom_style:
			attributes['padding'] = custom_style['padding']

		if 'borders' in custom_style:
			attributes['borders'] = custom_style['borders']

		if 'backgroundcolor' in custom_style:
			attributes['backgroundcolor'] = custom_style['backgroundcolor']

		if 'verticalalign' in custom_style:
			attributes['verticalalign'] = custom_style['verticalalign']

		if 'textalign' in custom_style:
			attributes['textalign'] = custom_style['textalign']

		if 'angle' in custom_style:
			attributes['angle'] = custom_style['angle']

	# borders
	if it_is_a_table_cell:
		if 'padding' in attributes:
			set_cell_padding(container, padding=attributes['padding'])

		if 'borders' in attributes:
			set_cell_border(container, borders=attributes['borders'])

		if 'backgroundcolor' in attributes:
			set_cell_bgcolor(container, color=attributes['backgroundcolor'])

		if 'verticalalign' in attributes:
			container.vertical_alignment = attributes['verticalalign']

		if 'textalign' in attributes:
			container.paragraphs[0].alignment = attributes['textalign']

		# text rotation, only for table._cell
		if 'angle' in attributes:
			rotate_text(cell=container, direction=attributes['angle'])

	else:
		if 'borders' in attributes:
			set_paragraph_border(element=container._p, borders=attributes['borders'])

		if 'backgroundcolor' in attributes:
			set_paragraph_bgcolor(element=container._element, color=attributes['backgroundcolor'])

		if 'textalign' in attributes:
			container.alignment = attributes['textalign']


''' whether the container is a document or not
'''
def is_document(container, nesting_level=0):
	# if container is n instance of table-cell
	if type(container) is document.Document:
		return True
	else:
		return False


''' whether the container is a table-cell or not
'''
def is_table_cell(container, nesting_level=0):
	# if container is n instance of table-cell
	if type(container) is table._Cell:
		return True
	else:
		return False


''' whether the container is a paragraph or not
'''
def is_paragraph(container, nesting_level=0):
	# if container is n instance of table-cell
	if isinstance(container, Paragraph):
		return True
	else:
		return False



# --------------------------------------------------------------------------------------------------------------------------------------------
# pictures, background image

''' insert image into a container
'''
def insert_image(container, inline_image, bookmark={}, nesting_level=0):
	if is_table_cell(container):
		run = container.paragraphs[0].add_run()

		# bookmark
		if bookmark:
			for k, v in bookmark.items():
				add_bookmark(paragraph=container.paragraphs[0], bookmark_name=k, bookmark_text=v)

		run.add_picture(inline_image.file_path, height=Inches(inline_image.image_height), width=Inches(inline_image.image_width))

		return container
	
	else:
		if is_document(container):
			paragraph = container.add_paragraph()

		elif is_paragraph(container):
			paragraph = container

		# bookmark
		if bookmark:
			for k, v in bookmark.items():
				add_bookmark(paragraph=paragraph, bookmark_name=k, bookmark_text=v)

		run = paragraph.add_run()
		if inline_image.image_width and inline_image.image_height:
			run.add_picture(inline_image.file_path, height=Inches(inline_image.image_height), width=Inches(inline_image.image_width))
		elif inline_image.image_width is None:
			run.add_picture(inline_image.file_path, height=Inches(inline_image.image_height))
		elif inline_image.image_height is None:
			run.add_picture(inline_image.file_path, width=Inches(inline_image.image_width))

		return paragraph


''' insert anchored image in a cell with margins
'''
def insert_cell_image(cell, inline_image, nesting_level=0):
    relative_h = 'column'
    relative_v = 'paragraph'

    # Optional offsets (in EMU)
    offX = pt_to_emu(inline_image.offsets_pt[0])
    offY = pt_to_emu(inline_image.offsets_pt[1])

    # If margin is a single number, apply to all sides. 
    if isinstance(inline_image.margin_pt, int):
        m = {'t': pt_to_emu(inline_image.margin_pt), 'b': pt_to_emu(inline_image.margin_pt), 'l': pt_to_emu(inline_image.margin_pt), 'r': pt_to_emu(inline_image.margin_pt)}

    # If it's a dict, use specific values.
    else:
        m = {'t': pt_to_emu(inline_image.margin_pt.get('t', 0)), 'b': pt_to_emu(inline_image.margin_pt.get('b', 0)), 
             'l': pt_to_emu(inline_image.margin_pt.get('l', 0)), 'r': pt_to_emu(inline_image.margin_pt.get('r', 0))}

    mapping = {
        'left top':      {'h': 'left',   'v': 'top'},
        'center top':    {'h': 'center', 'v': 'top'},
        'right top':     {'h': 'right',  'v': 'top'},
        'left middle':   {'h': 'left',   'v': 'center'},
        'center middle': {'h': 'center', 'v': 'center'},
        'right middle':  {'h': 'right',  'v': 'center'},
        'left bottom':   {'h': 'left',   'v': 'bottom'},
        'center bottom': {'h': 'center', 'v': 'bottom'},
        'right bottom':  {'h': 'right',  'v': 'bottom'},
    }
    
    aligns = mapping.get(inline_image.position, mapping['center middle'])
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    shape = run.add_picture(inline_image.file_path, width=Inches(inline_image.image_width))
    inline = shape._inline
    rId = inline.graphic.graphicData.pic.blipFill.blip.embed
    cx, cy = inline.extent.cx, inline.extent.cy
    pic_id = random.randint(1000, 10000)
    name = f"Picture-{pic_id}"

    # Logic: distT, distB, distL, distR control the "text wrap buffer"
    anchor_xml = f'''
	<wp:anchor distT="{m['t']}" distB="{m['b']}" distL="{m['l']}" distR="{m['r']}"
            simplePos="0" relativeHeight="251658240" behindDoc="0"
            locked="0" layoutInCell="1" allowOverlap="1"
            {nsdecls('wp', 'a', 'pic', 'r')}>
		<wp:simplePos x="0" y="0" />
        <wp:positionH relativeFrom="{relative_h}">
            <wp:align>{aligns['h']}</wp:align>
        </wp:positionH>
        <wp:positionV relativeFrom="{relative_v}">
            <wp:align>{aligns['v']}</wp:align>
        </wp:positionV>
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
		<wp:wrapSquare wrapText="bothSides" />
		<wp:docPr id="{pic_id}" name="{name}" />
		<wp:cNvGraphicFramePr>
			<a:graphicFrameLocks
				xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
				noChangeAspect="1" />
		</wp:cNvGraphicFramePr>
		<a:graphic
			xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
			<a:graphicData
				uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
				<pic:pic
					xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
					<pic:nvPicPr>
						<pic:cNvPr id="{pic_id}" name="{name}" />
						<pic:cNvPicPr />
					</pic:nvPicPr>
					<pic:blipFill>
						<a:blip r:embed="{rId}">
						</a:blip>
						<a:stretch>
							<a:fillRect />
						</a:stretch>
					</pic:blipFill>
					<pic:spPr>
						<a:xfrm>
							<a:off x="0" y="0" />
							<a:ext cx="{cx}" cy="{cy}" />
						</a:xfrm>
						<a:prstGeom prst="rect">
							<a:avLst />
						</a:prstGeom>
					</pic:spPr>
				</pic:pic>
			</a:graphicData>
		</a:graphic>
	</wp:anchor>
    '''

    anchor = parse_xml(anchor_xml)
    inline.getparent().replace(inline, anchor)


''' create page background
	<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658240" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
		<wp:simplePos x="0" y="0" />
		<wp:positionH relativeFrom="page">
			<wp:posOffset>0</wp:posOffset>
		</wp:positionH>
		<wp:positionV relativeFrom="page">
			<wp:posOffset>0</wp:posOffset>
		</wp:positionV>
		<wp:extent cx="7562000" cy="10689336" />
		<wp:effectExtent l="0" t="0" r="0" b="0" />
		<wp:wrapNone />
		<wp:docPr id="1" name="Picture 1" />
		<wp:cNvGraphicFramePr>
			<a:graphicFrameLocks
				xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" />
		</wp:cNvGraphicFramePr>
		<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
			<a:graphicData
				uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
				<pic:pic
					xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
					<pic:nvPicPr>
						<pic:cNvPr id="0" name="toll-booth.jpg" />
						<pic:cNvPicPr />
					</pic:nvPicPr>
					<pic:blipFill>
						<a:blip r:embed="rId8">
							<a:extLst>
								<a:ext uri="{28A0092B-C50C-407E-A947-70E740481C1C}">
									<a14:useLocalDpi
										xmlns:a14="http://schemas.microsoft.com/office/drawing/2010/main"
										val="0" />
								</a:ext>
							</a:extLst>
						</a:blip>
						<a:stretch>
							<a:fillRect />
						</a:stretch>
					</pic:blipFill>
					<pic:spPr>
						<a:xfrm>
							<a:off x="0" y="0" />
							<a:ext cx="7562000" cy="10692000" />
						</a:xfrm>
						<a:prstGeom prst="rect">
							<a:avLst />
						</a:prstGeom>
					</pic:spPr>
				</pic:pic>
			</a:graphicData>
		</a:graphic>
		<wp14:sizeRelH relativeFrom="margin">
			<wp14:pctWidth>0</wp14:pctWidth>
		</wp14:sizeRelH>
		<wp14:sizeRelV relativeFrom="margin">
			<wp14:pctHeight>0</wp14:pctHeight>
		</wp14:sizeRelV>
	</wp:anchor>
'''
def create_page_background(docx, header, background_image_path, page_width_inches, page_height_inches, nesting_level=0):
	drawing_xml = '''
	<w:drawing>
		<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
			<wp:simplePos x="0" y="0" />
			<wp:positionH relativeFrom="page">
				<wp:posOffset>0</wp:posOffset>
			</wp:positionH>
			<wp:positionV relativeFrom="page">
				<wp:posOffset>0</wp:posOffset>
			</wp:positionV>
			<wp:extent cx="{cx}" cy="{cy}" />
			<wp:effectExtent l="0" t="0" r="0" b="0" />
			<wp:wrapNone />
			<wp:docPr id="{doc_id}" name="{doc_name}" />
			<wp:cNvGraphicFramePr>
				<a:graphicFrameLocks
					xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" />
			</wp:cNvGraphicFramePr>
			<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
				<a:graphicData
					uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
					<pic:pic
						xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
						<pic:nvPicPr>
							<pic:cNvPr id="{image_id}" name="{image_name}" />
							<pic:cNvPicPr />
						</pic:nvPicPr>
						<pic:blipFill>
							<a:blip r:embed="{rid}">
							</a:blip>
							<a:stretch>
								<a:fillRect />
							</a:stretch>
						</pic:blipFill>
						<pic:spPr>
							<a:xfrm>
								<a:off x="0" y="0" />
								<a:ext cx="{cx}" cy="{cy}" />
							</a:xfrm>
							<a:prstGeom prst="rect">
								<a:avLst />
							</a:prstGeom>
						</pic:spPr>
					</pic:pic>
				</a:graphicData>
			</a:graphic>
			<wp14:sizeRelH relativeFrom="margin">
				<wp14:pctWidth>0</wp14:pctWidth>
			</wp14:sizeRelH>
			<wp14:sizeRelV relativeFrom="margin">
				<wp14:pctHeight>0</wp14:pctHeight>
			</wp14:sizeRelV>
		</wp:anchor>
	</w:drawing>
	'''

	# get the current/last paragraph
	first_para = header.paragraphs[0]
	first_run = first_para.add_run()

	# add a new paragraph paragraph
	new_para = docx.add_paragraph()

	# create a run
	new_run = new_para.add_run()

	# put the image
	shape = new_run.add_picture(image_path_or_stream=background_image_path, height=Inches(page_height_inches))
	# current_drawing_element = new_run._r.xpath('//w:drawing')[0]

	try:
		# tweak the generated inline image
		parser = etree.XMLParser(recover=True)

		cx = str(emu(page_width_inches))
		cy = str(emu(page_height_inches))


		docPr = new_run._r.xpath('.//wp:docPr')[0]
		doc_id = docPr.get('id')
		doc_name = docPr.get('name')

		cNvPr = new_run._r.xpath('.//pic:cNvPr')[0]
		image_id = cNvPr.get('id')
		image_name = cNvPr.get('name')
		trace(f"bg image [{image_name}]", nesting_level=nesting_level)

		blip = new_run._r.xpath('.//a:blip')[0]

		rid = None
		try:
			rid = blip.xpath('./@r:embed')[0]
			# trace(f"rid [{rid}]")
		except:
			warn(f"rid not found under blip")
		
		if rid is None:
			try:
				# Assuming blip_element is your element object
				# namespaces = {'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}

				# Access via the attrib dictionary
				# Note: You must use the {URI} format as the key
				# attr_key = f"{{{namespaces['r']}}}embed"
				attr_key = f"r:embed"
				rid = blip.attrib.get(attr_key)
				trace(f"rid [{rid}]")

			except Exception as e:
				warn(f"r:embed not found under blip")
				raise e

		new_drawing_element = etree.XML(drawing_xml.format(cx=cx, cy=cy, doc_id=doc_id, doc_name=doc_name, image_id=image_id, image_name=image_name, rid=rid), parser)
		# put the new drawing into first-run
		first_run._r.append(new_drawing_element)

	except Exception as e:
		error(traceback.format_exc())

	# remove the new-para
	delete_paragraph(new_para)


''' Inserts an image into the header, stretches it to page size, and makes it floating behind text.
'''
def add_background_image_to_header(docx_section, image_path, width, height, nesting_level=0):
    # Put it in its own paragraph at start
	for header in [docx_section.header, docx_section.even_page_header]:
		p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
		run = p.add_run()
		inline = run.add_picture(image_path, width=width, height=height)
		inline_to_anchored_behind(inline, width=width, height=height)
		unlock_picture_aspect_ratio(inline)
		force_picture_transform(inline, width_in=width, height_in=height)


''' insert an image as cell background
'''
def create_cell_background(cell, image_path, width, height, nesting_level=0):
	
	# 1. Add a run to the paagraph
	paragraph = cell.paragraphs[0]
	run = paragraph.add_run()

	# 2. Add the picture (initially inline)
	picture = run.add_picture(image_path, width=Inches(width), height=Inches(height))

	# 3. Get the XML element and change it from 'inline' to 'anchor'
	inline = picture._inline

	cx = str(emu(width))
	cy = str(emu(height))

	anchor_xml = f"""
	<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" 
			   relativeHeight="0" behindDoc="1" locked="0" layoutInCell="1" 
			   allowOverlap="1" {nsdecls('wp', 'a', 'pic', 'r')}>
		<wp:simplePos x="0" y="0"/>
		<wp:positionH relativeFrom="column">
			<wp:posOffset>0</wp:posOffset>
		</wp:positionH>
		<wp:positionV relativeFrom="line">
			<wp:posOffset>0</wp:posOffset>
		</wp:positionV>
		<wp:extent cx="{cx}" cy="{cy}"/>
		<wp:effectExtent l="0" t="0" r="0" b="0"/>
		<wp:wrapNone/>
		<wp:docPr id="1" name="CellBackground"/>
		<wp:cNvGraphicFramePr/>
	</wp:anchor>
	"""

	# parser = etree.XMLParser(recover=False)
	# anchor = etree.XML(anchor_xml.format(cx=cx, cy=cy), parser)
	anchor = parse_xml(anchor_xml)

	# 4. Transfer the graphic data from the inline tag to the anchor tag
	# graphic = inline.xpath('.//a:graphic', namespaces=inline.nsmap)[0]
	graphic = inline.xpath('.//a:graphic')[0]
	anchor.append(graphic)
	
	# 5. Replace the original inline XML with our new anchor XML
	inline.getparent().replace(inline, anchor)

	# move this run to the 0th index of the paragraph
	move_run_to_start(paragraph, run)

	zero_paragraph_spacing(paragraph)


''' insert an image as cell background
'''
def create_cell_background_new(cell, image_path, width, height, nesting_level=0):

	# 1. Clear outer cell and apply vertical alignment
    cell._element.clear_content()

    # 2. Create a stacking table (1x1)
    stack_table = cell.add_table(rows=1, cols=1)
    remove_table_borders(stack_table)
    stack_cell = stack_table.cell(0, 0)
    remove_cell_margins(stack_cell)

    # 3. Background layer (image)
    bg_table = stack_cell.add_table(rows=1, cols=1)
    remove_table_borders(bg_table)
    bg_cell = bg_table.cell(0, 0)
    remove_cell_margins(bg_cell)

    p_img = bg_cell.paragraphs[0]
    run = p_img.add_run()
    run.add_picture(image_path, width=width, height=height)
    zero_paragraph_spacing(p_img)

    # 4. Overlay layer (text)
    fg_table = stack_cell.add_table(rows=1, cols=1)
    remove_table_borders(fg_table)
    fg_cell = fg_table.cell(0, 0)
    remove_cell_margins(fg_cell)

    p_txt = fg_cell.paragraphs[0]
    zero_paragraph_spacing(p_txt)

    return fg_cell



# --------------------------------------------------------------------------------------------------------------------------------------------
# table and table-cell

''' create a Table
'''
def create_table(container, num_rows, num_cols, container_width=None, nesting_level=0):
	tbl = None

	# debug(f"... creating ({num_rows} x {num_cols} : width = {container_width}) table for {type(container)}")
	if type(container) is section._Header or type(container) is section._Footer:
		# if the conrainer is a Header/Footer
		tbl = container.add_table(num_rows, num_cols, container_width)
	elif is_table_cell(container):
		# if the conrainer is a Cell
		tbl = container.add_table(num_rows, num_cols)
	elif is_document(container):
		# if the conrainer is a Document
		tbl = container.add_table(num_rows, num_cols)

	tbl.style = 'PlainTable'
	tbl.autofit = False

	return tbl


''' polish a table
'''
def polish_table(table, nesting_level=0):
	for r in table.rows:
		# no preferred width for the last column
		c = r._tr.tc_lst[-1]
		#for c in r._tr.tc_lst:
		tcW = c.tcPr.tcW
		tcW.type = 'auto'
		tcW.w = 0


''' remove table borders
'''
def remove_table_borders(table, nesting_level=0):
    tbl = table._element
    tblPr = tbl.tblPr or OxmlElement("w:tblPr")

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(ns.qn("w:val"), "nil")
        borders.append(el)

    tblPr.append(borders)
    tbl.append(tblPr)


''' set repeat table row on every new page
'''
def set_repeat_table_header(row, nesting_level=0):
	tr = row._tr
	trPr = tr.get_or_add_trPr()
	tblHeader = OxmlElement('w:tblHeader')
	tblHeader.set(qn('w:val'), "true")
	trPr.append(tblHeader)
	return row


''' set table-cell borders
'''
def set_cell_border(cell: table._Cell, borders, nesting_level=0):
	tc = cell._tc
	tcPr = tc.get_or_add_tcPr()

	# check for tag existnace, if none found, then create one
	tcBorders = tcPr.first_child_found_in("w:tcBorders")
	if tcBorders is None:
		tcBorders = OxmlElement('w:tcBorders')
		tcPr.append(tcBorders)

	# list over all available tags
	for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
		edge_data = borders.get(edge)
		if edge_data:
			tag = 'w:{}'.format(edge)

			# check for tag existnace, if none found, then create one
			element = tcBorders.find(qn(tag))
			if element is None:
				element = OxmlElement(tag)
				tcBorders.append(element)

			# looks like order of attributes is important
			for key in ["sz", "val", "color", "space", "shadow"]:
				if key in edge_data:
					element.set(qn('w:{}'.format(key)), str(edge_data[key]))


''' set table-cell bgcolor
'''
def set_cell_bgcolor(cell: table._Cell, color, nesting_level=0):
	xml = r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color)
	shading_elm_1 = parse_xml(xml)
	cell._tc.get_or_add_tcPr().append(shading_elm_1)


''' set table-cell padding - space between border and text
'''
def set_cell_padding(cell: table._Cell, padding, nesting_level=0):
	tc = cell._tc
	tcPr = tc.get_or_add_tcPr()
	tcMar = OxmlElement('w:tcMar')

	for m in ["top", "start", "bottom", "end", ]:
		if m in padding:
			node = OxmlElement("w:{}".format(m))
			node.set(qn('w:w'), str(padding.get(m)))
			node.set(qn('w:type'), 'dxa')
			tcMar.append(node)

	tcPr.append(tcMar)


''' remove cell margins
'''
def remove_cell_margins(cell, nesting_level=0):
    tcPr = cell._element.tcPr or OxmlElement("w:tcPr")

    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(ns.qn("w:w"), "0")
        el.set(ns.qn("w:type"), "dxa")
        tcMar.append(el)

    tcPr.append(tcMar)
    cell._element.append(tcPr)


''' rotate text
'''
def rotate_text(cell: table._Cell, direction: str, nesting_level=0):
	# direction: tbRl -- top to bottom, btLr -- bottom to top
	assert direction in ("tbRl", "btLr")
	tc = cell._tc
	tcPr = tc.get_or_add_tcPr()
	textDirection = OxmlElement('w:textDirection')
	textDirection.set(qn('w:val'), direction)  # btLr tbRl
	tcPr.append(textDirection)


''' copy cell border from one cell to another
'''
def copy_cell_border(from_cell: table._Cell, to_cell: table._Cell, nesting_level=0):
	from_tc = from_cell._tc
	from_tcPr = from_tc.get_or_add_tcPr()

	to_tc = to_cell._tc
	to_tcPr = to_tc.get_or_add_tcPr()

	from_tcBorders = from_tcPr.first_child_found_in("w:tcBorders")
	to_tcBorders = to_tcPr.first_child_found_in("w:tcBorders")
	if from_tcBorders is not None:
		if to_tcBorders is None:
			to_tcBorders = deepcopy(from_tcBorders)
			to_tc.get_or_add_tcPr().append(to_tcBorders)


''' assert if a cell is not a covered cell
'''
def assert_not_merge_continuation(cell, nesting_level=0):
    tcPr = cell._element.tcPr
    if tcPr is None:
        return
    vMerge = tcPr.find(ns.qn("w:vMerge"))
    if vMerge is not None and vMerge.get(ns.qn("w:val")) == "continue":
        raise ValueError(
            "Cannot insert content into a merged-cell continuation. "
            "Use the merge-restart cell only."
        )


# --------------------------------------------------------------------------------------------------------------------------------------------
# paragraphs and texts

''' write a paragraph in a given style
'''
def create_paragraph(docx, container, text_content=None, run_list=None, paragraph_attributes=None, text_attributes=None, bg_image=None, footnote_list={}, bookmark={}, directives=True, nesting_level=0):
	# create or get the paragraph
	if type(container) is section._Header or type(container) is section._Footer:
		# if the container is a Header/Footer
		paragraph = container.add_paragraph()

	elif is_table_cell(container):
		# if the conrainer is a Cell, the Cell already has an empty paragraph
		paragraph = container.paragraphs[0]
		if bg_image is not None:
			fg_cell = create_cell_background(cell=container, image_path=bg_image.file_path, width=bg_image.container_width, height=bg_image.container_height, nesting_level=nesting_level+1)
			paragraph = fg_cell.paragraphs[0]
			# pass

	elif is_document(container):
		# if the conrainer is a Document
		paragraph = container.add_paragraph()

	elif is_paragraph(container):
		pass

	else:
		# if the conrainer is anything else
		warn(f"container is neither document, nor paragraph, nor header/footer, nor table cell .. adding a paragraph", nesting_level=nesting_level)
		paragraph = container.add_paragraph()


	# process inline blocks like footnotes, latex inside the content
	if run_list is not None:
		# run lists are a series of runs inside the paragraph
		for text_run in run_list:
			process_inline_blocks(docx=docx, paragraph=paragraph, text_content=text_run['text'], text_attributes=text_run['text-attributes'], footnote_list=footnote_list, nesting_level=nesting_level+1)

	elif text_content is not None:
		process_inline_blocks(docx=docx, paragraph=paragraph, text_content=text_content, text_attributes=text_attributes, footnote_list=footnote_list, nesting_level=nesting_level+1)


	# bookmark
	if bookmark:
		for k, v in bookmark.items():
			# debug(f"creating bookmark {k} : {v}")
			add_bookmark(paragraph=paragraph, bookmark_name=k, bookmark_text=v, nesting_level=nesting_level+1)

	# apply the style if any
	apply_paragraph_attributes(docx=docx, paragraph=paragraph, paragraph_attributes=paragraph_attributes, nesting_level=nesting_level+1)

	return paragraph


''' set paragraph borders
	{
		"top":		{"sz": , "val": , "color": , "space": , "shadow": },
		"start":	{},
		"bottom":	{},
		"end": 		{},
	}
	size is in 1/8 pt units
	val is any of dotted/dashed/single/thick/triple/double/none
	space is space/padding between text and border in pt
'''
def set_paragraph_border(element, borders, nesting_level=0):
	pPr = element.get_or_add_pPr()

	# check for tag existnace, if none found, then create one
	# pBorders = OxmlElement('w:pBorders')
	pBorders = pPr.first_child_found_in("w:pBorders")
	if pBorders is None:
		pBorders = OxmlElement('w:pBorders')

	# list over all available tags
	for edge in ('top', 'start', 'bottom', 'end'):
		edge_data = borders.get(edge)
		if edge_data:
			edge_str = edge
			if edge_str == 'start': edge_str = 'left'
			if edge_str == 'end': edge_str = 'right'

			border = OxmlElement(f'w:{edge_str}')
			border.set(qn('w:val'), str(edge_data['val']))     # single, double, dashed, dotted
			border.set(qn('w:sz'), str(edge_data['sz']))          # border width (1/8 pt units)
			border.set(qn('w:space'), str(edge_data['space']))        # space between text and border
			border.set(qn('w:color'), str(edge_data['color']))   # hex color
			pBorders.append(border)

	pPr.append(pBorders)


''' set paragraph bgcolor
'''
def set_paragraph_bgcolor(element, color, nesting_level=0):
	shading = OxmlElement('w:shd')

	# required
	shading.set(qn('w:val'), 'clear')
	# text color
	shading.set(qn('w:color'), 'auto')
	# background color (hex, no #)	
	shading.set(qn('w:fill'), color.lstrip('#'))

	element.get_or_add_pPr().append(shading)


''' apply paragraph attributes to a paragraph
'''
def	apply_paragraph_attributes(docx, paragraph, paragraph_attributes, nesting_level=0):
	# apply the style if any
	if paragraph_attributes:
		# apply paragraph attrubutes
		if 'stylename' in paragraph_attributes:
			style_name = paragraph_attributes['stylename']
			# check if the style exists
			style = get_style_by_name(docx=docx, style_name=style_name)
			if style is not None:
				paragraph.style = style_name
			else:
				# trace(f"style {style_name} is not a standard style", nesting_level=nesting_level)
				pass

		pf = paragraph.paragraph_format

		# process new-page
		if 'breakbefore' in paragraph_attributes:
			pf.page_break_before = True

		# process keep-with-next
		if 'keepwithnext' in paragraph_attributes:
			pf.keep_with_next = True

		# process keep-with-previous
		if 'keepwithprevious' in paragraph_attributes:
			# docx does not support keep_with_previous
			# pf.keep_with_previous = True
			pass

	return paragraph


''' remove a paragraph
'''
def delete_paragraph(paragraph, nesting_level=0):
	p = paragraph._element
	p.getparent().remove(p)
	# p._p = p._element = None


''' tex/character style for text run
'''
def set_text_style(run, text_attributes, nesting_level=0):
	if text_attributes:
		if 'fontweight' in text_attributes:
			run.bold = True
		else:
			run.bold = False

		if 'fontstyle' in text_attributes:
			run.italic = True
		else:
			run.italic = False

		if 'textlinethroughstyle' in text_attributes:
			run.strike = True
		else:
			run.strike = False

		if 'textunderlinestyle' in text_attributes:
			run.underline = True
		else:
			run.underline = False

		run.font.name = text_attributes['fontname']
		run.font.size = Pt(text_attributes['fontsize'])

		fgcolor = text_attributes.get('color')
		run.font.color.rgb = RGBColor(fgcolor.red, fgcolor.green, fgcolor.blue)


''' inside a paragraph move a run from anywhere to start as the 0th run
'''
def move_run_to_start(paragraph, run_obj, nesting_level=0):
	p_element = paragraph._p
	r_element = run_obj._element
    
	p_element.remove(r_element)
    
    # Check if properties exist; if so, put run after them. 
	pPr = p_element.pPr
	if pPr is not None:
		pPr.addnext(r_element)

    # If not, put run at the very top.
	else:
		p_element.insert(0, r_element)
		

''' remove all spacing around a paragraph
'''
def zero_paragraph_spacing(paragraph, nesting_level=0):
    pf = paragraph.paragraph_format
    pf.space_before = 0
    pf.space_after = 0



# --------------------------------------------------------------------------------------------------------------------------------------------
# footnotes, bookmarks, links, latex, mathml related

''' process inline blocks inside a text and add to a paragraph
'''
def process_inline_blocks(docx, paragraph, text_content, text_attributes, footnote_list, nesting_level=0):
	# process FN{...} first, we get a list of block dicts
	inline_blocks = process_footnotes(
		text_content=text_content, footnote_list=footnote_list
	)

	# process LATEX{...} for each text item
	new_inline_blocks = []
	for inline_block in inline_blocks:
		# process only 'text'
		if "text" in inline_block:
			new_inline_blocks = new_inline_blocks + process_latex_blocks(inline_block["text"])

		else:
			new_inline_blocks.append(inline_block)

	inline_blocks = new_inline_blocks

	# process PAGE{..} for each text item
	new_inline_blocks = []
	for inline_block in inline_blocks:
		# process only 'text'
		if "text" in inline_block:
			new_inline_blocks = new_inline_blocks + process_bookmark_page_blocks(inline_block["text"])

		else:
			new_inline_blocks.append(inline_block)

	inline_blocks = new_inline_blocks

	# process LINK{..} for each text item
	new_inline_blocks = []
	for inline_block in inline_blocks:
		# process only 'text'
		if "text" in inline_block:
			new_inline_blocks = new_inline_blocks + process_links(inline_block["text"])

		else:
			new_inline_blocks.append(inline_block)

	inline_blocks = new_inline_blocks

	# we are ready to prepare the content
	for inline_block in inline_blocks:
		if "text" in inline_block:
			run = paragraph.add_run(inline_block["text"])
			set_text_style(run=run, text_attributes=text_attributes)

		elif "fn" in inline_block:
			create_footnote_bayoo(paragraph=paragraph, footnote_tuple=inline_block["fn"])
			# create_footnote(docx=docx, paragraph=paragraph, footnote_tuple=inline_block["fn"])

		elif "latex" in inline_block:
			run = paragraph.add_run()
			create_latex(container=run, latex_content=inline_block["latex"])
			set_text_style(run=run, text_attributes=text_attributes)

		elif "page-num" in inline_block:
			run = add_page_reference(paragraph=paragraph, bookmark_name='')
			set_text_style(run=run, text_attributes=text_attributes)

		elif "page-count" in inline_block:
			run = add_page_reference(paragraph=paragraph, bookmark_name='*')
			set_text_style(run=run, text_attributes=text_attributes)

		elif "bookmark-page" in inline_block:
			# add_link(paragraph=paragraph, link_to=inline_block["page"].strip(), text=inline_block["page"].strip(), tool_tip=None)
			run = add_page_reference(paragraph=paragraph, bookmark_name=inline_block["bookmark-page"].strip())
			set_text_style(run=run, text_attributes=text_attributes)

		elif 'link' in inline_block:
			target, anchor = inline_block['link'][0], inline_block['link'][1]
			create_hyperlink(attach_to=paragraph, anchor=anchor, target=target)


''' process footnotes inside text
'''
def process_footnotes(text_content, footnote_list, nesting_level=0):
	# if text contains footnotes we make a list containing texts->footnote->text->footnote ......
	texts_and_footnotes = []

	# find out if there is any match with FN#key inside the text_content
	pattern = r'FN{[^}]+}'
	current_index = 0
	for match in re.finditer(pattern, text_content):
		footnote_key = match.group()[3:-1]
		if footnote_key in footnote_list:
			# debug(f".... footnote {footnote_key} found at {match.span()} with description")
			# we have found a footnote, we add the preceding text and the footnote spec into the list
			footnote_start_index, footnote_end_index = match.span()[0], match.span()[1]
			if footnote_start_index >= current_index:
				# there are preceding text before the footnote
				texts_and_footnotes.append({'text': text_content[current_index:footnote_start_index]})
				texts_and_footnotes.append({'fn': (footnote_key, footnote_list[footnote_key])})
				current_index = footnote_end_index
		else:
			warn(f".... footnote {footnote_key} found at {match.span()}, but no details found")
			# this is not a footnote, ignore it
			footnote_start_index, footnote_end_index = match.span()[0], match.span()[1]
			# current_index = footnote_end_index + 1

	# there may be trailing text
	texts_and_footnotes.append({'text': text_content[current_index:]})

	return texts_and_footnotes


''' process latex blocks inside text
'''
def process_latex_blocks(text_content, nesting_level=0):
	# find out if there is any match with LATEX$...$ inside the text_content
	texts_and_latex = []

	pattern = r'LATEX\$[^$]+\$'
	current_index = 0
	for match in re.finditer(pattern, text_content):
		latex_content = match.group()[6:-1]

		# we have found a latex block, we add the preceding text and the latex block into the list
		latex_start_index, latex_end_index = match.span()[0], match.span()[1]
		if latex_start_index >= current_index:
			# there are preceding text before the latex
			text = text_content[current_index:latex_start_index]

			texts_and_latex.append({'text': text})

			texts_and_latex.append({'latex': latex_content})

			current_index = latex_end_index


	# there may be trailing text
	text = text_content[current_index:]

	texts_and_latex.append({'text': text})

	return texts_and_latex


''' process bookmark page ref inside text
'''
def process_bookmark_page_blocks(text_content, nesting_level=0):
	# find out if there is any match with PAGE{...} inside the text_content
	texts_and_bookmarks = []

	pattern = r'PAGE{[^}]*}'
	current_index = 0
	for match in re.finditer(pattern, text_content):
		bookmark_content = match.group()[5:-1].strip()

		# we have found a PAGE block, we add the preceding text and the PAGE block into the list
		bookmark_start_index, bookmark_end_index = match.span()[0], match.span()[1]
		if bookmark_start_index >= current_index:
			# there are preceding text before the PAGE
			text = text_content[current_index:bookmark_start_index]

			texts_and_bookmarks.append({'text': text})

			# PAGE{} means current page, PAGE{*} means number of pages, PAGE{XYZ} means page number where bookmark XYZ is set
			if bookmark_content == '':
				texts_and_bookmarks.append({"page-num": None})

			elif bookmark_content == '*':
				texts_and_bookmarks.append({'page-count': None})

			else:
				texts_and_bookmarks.append({'bookmark-page': bookmark_content})

			current_index = bookmark_end_index


	# there may be trailing text
	text = text_content[current_index:]

	texts_and_bookmarks.append({'text': text})

	return texts_and_bookmarks


''' process external url or bookmark links
'''
def process_links(text_content, nesting_level=0):
	# find out if there is any match with LINK{...}{} inside the text_content
	links = []

	# LINK patterns are like LINK{target}{text} or LINK{target}
	pattern = r'LINK({[^}]*}){1,2}'
	current_index = 0
	for match in re.finditer(pattern, text_content):
		link_content_pattern = r'([^{}]+)'
		i = 0
		target, anchor = None, None
		for content_match in re.finditer(link_content_pattern, match.group()):
			if i == 1:
				target = content_match.group()
			elif i == 2:
				anchor = content_match.group()

			i = i + 1

		# we have found a LINK block, we add the preceding text and the LINK block into the list
		link_start_index, link_end_index = match.span()[0], match.span()[1]
		if link_start_index >= current_index:
			# there are preceding text before the link
			text = text_content[current_index:link_start_index]
			links.append({'text': text})

			# LINK patterns are like LINK{target}{text} or LINK{target}
			if target:
				links.append({"link": [target, anchor]})

			current_index = link_end_index

	# there may be trailing text
	text = text_content[current_index:]

	links.append({'text': text})

	return links


''' create a footnote
'''
def create_footnote(docx, paragraph, footnote_tuple, nesting_level=0):
	# create footnote reference
	r = paragraph.add_run()._r
	fn_ref = OxmlElement('w:footnoteReference')
	fn_ref.set(qn('w:id'), footnote_tuple[0])
	r.append(fn_ref)

	# access footnotes part
	footnotes_part = docx.part._footnotes_part

	# create footnote
	fn = OxmlElement('w:footnote')
	fn.set(qn('w:id'), footnote_tuple[0])

	p_fn = OxmlElement('w:p')
	r_fn = OxmlElement('w:r')
	t_fn = OxmlElement('w:t')
	t_fn.text = footnote_tuple[1]

	r_fn.append(t_fn)
	p_fn.append(r_fn)
	fn.append(p_fn)

	footnotes_part._element.append(fn)


''' create a footnote
'''
def create_footnote_bayoo(paragraph, footnote_tuple, nesting_level=0):
	paragraph.add_footnote(footnote_tuple[1])


''' add a latex/mathml run into a paragraph
'''
def create_latex(container, latex_content, nesting_level=0):
	if latex_content is not None:
		mathml_output = latex2mathml.converter.convert(strip_math_mode_delimeters(latex_content))

		# Convert MathML (MML) into Office MathML (OMML) using a XSLT stylesheet
		tree = etree.fromstring(mathml_output)
		xslt = etree.parse(mml2omml_stylesheet_path)

		transform = etree.XSLT(xslt)
		new_dom = transform(tree)

		container._element.append(new_dom.getroot())


''' create a bookmark
'''
def add_bookmark(paragraph, bookmark_name, bookmark_text='', nesting_level=0):
	# debug(f"creating bookmark {bookmark_name} : {bookmark_text}")
	start = OxmlElement('w:bookmarkStart')
	start.set(qn('w:id'), '0')
	start.set(qn('w:name'), bookmark_name)

	# text = OxmlElement('w:r')
	# text.text = bookmark_text

	end = OxmlElement('w:bookmarkEnd')
	end.set(qn('w:id'), '0')
	end.set(qn('w:name'), bookmark_name)

	paragraph._p.insert(0, start)
	# paragraph._p.append(text)
	paragraph._p.append(end)


''' add a PAGE refernce 
'''
def add_page_reference(paragraph, bookmark_name, nesting_level=0):
	run = paragraph.add_run()

	# create a new element and set attributes
	fldCharBegin = OxmlElement('w:fldChar')
	fldCharBegin.set(qn('w:fldCharType'), 'begin')

	# actual PAGE reference
	instrText = OxmlElement("w:instrText")
	instrText.set(qn("xml:space"), "preserve")

	if bookmark_name == '':
		instrText.text = 'PAGE \\* MERGEFORMAT'

	elif bookmark_name == '*':
		instrText.text = 'NUMPAGES \\* MERGEFORMAT'

	else:
		instrText.text = f"PAGEREF {bookmark_name} \\h"

	fldCharEnd = OxmlElement('w:fldChar')
	fldCharEnd.set(qn('w:fldCharType'), 'end')

	r_element = run._r
	r_element.append(fldCharBegin)
	r_element.append(instrText)
	r_element.append(fldCharEnd)

	# p_element = paragraph._p

	return run


''' add link for bookmarks
'''
def add_link(paragraph, link_to, text, tool_tip=None, nesting_level=0):
	# create hyperlink node
	hyperlink = OxmlElement('w:hyperlink')

	# set attribute for link to bookmark
	hyperlink.set(qn('w:anchor'), link_to,)

	if tool_tip is not None:
		# set attribute for link to bookmark
		hyperlink.set(qn('w:tooltip'), tool_tip,)

	new_run = OxmlElement('w:r')
	rPr = OxmlElement('w:rPr')
	new_run.append(rPr)
	new_run.text = text
	hyperlink.append(new_run)
	r = paragraph.add_run()
	r._r.append(hyperlink)
	r.font.name = "Calibri"
	r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
	r.font.underline = True


''' This is needed if we are using the builtin style
'''
def get_or_create_hyperlink_style(d, nesting_level=0):
	"""If this document had no hyperlinks so far, the builtin
	   Hyperlink style will likely be missing and we need to add it.
	   There's no predefined value, different Word versions
	   define it differently.
	   This version is how Word 2019 defines it in the
	   default theme, excluding a theme reference.
	"""
	if "Hyperlink" not in d.styles:
		if "Default Character Font" not in d.styles:
			ds = d.styles.add_style("Default Character Font", WD_STYLE_TYPE.CHARACTER, True)
			ds.element.set(qn('w:default'), "1")
			ds.priority = 1
			ds.hidden = True
			ds.unhide_when_used = True
			del ds
		hs = d.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER, True)
		hs.base_style = d.styles["Default Character Font"]
		hs.unhide_when_used = True
		hs.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
		hs.font.underline = True
		del hs

	return "Hyperlink"


''' add the hyperlink to a run
'''
def add_hyperlink(paragraph, text, url, nesting_level=0):
	# This gets access to the document.xml.rels file and gets a new relation id value
	part = paragraph.part
	r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

	# Create the w:hyperlink tag and add needed values
	hyperlink = OxmlElement('w:hyperlink')
	hyperlink.set(qn('r:id'), r_id, )

	# Create a new run object (a wrapper over a 'w:r' element)
	new_run = paragraph.add_run(OxmlElement('w:r'))
	new_run.text = text

	# Set the run's style to the builtin hyperlink style, defining it if necessary
	new_run.style = get_or_create_hyperlink_style(part.document)
	# Alternatively, set the run's formatting explicitly
	# new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
	# new_run.font.underline = True

	# Join all the xml elements together
	hyperlink.append(new_run._element)
	paragraph._p.append(hyperlink)

	return hyperlink


''' create a hyperlink
'''
def create_hyperlink(attach_to, anchor, target, nesting_level=0):
	# if the anchor is not an url, it is a bookmark
	if not target.startswith('http'):
		target_text = target.strip()
		if anchor:
			anchor_text = anchor
		else:
			# TODO: it should actually be the text associated with the target bookmark
			anchor_text = target_text
	else:
		target_text = target.strip()
		if anchor:
			anchor_text = anchor
		else:
			anchor_text = target_text

	if not target.startswith('http'):
		# create hyperlink node
		hyperlink = OxmlElement('w:hyperlink')

		# set attribute for link to bookmark
		hyperlink.set(qn('w:anchor'), target_text,)
		hyperlink.set(qn('w:tooltip'), target_text,)

		# change the font color, and add underline
		rPr = OxmlElement('w:rPr')
		c = OxmlElement('w:color')
		c.set(qn('w:val'), '2A6099')
		rPr.append(c)
		u = OxmlElement('w:u')
		u.set(qn('w:val'), 'single')
		rPr.append(u)

		new_run = OxmlElement('w:r')
		new_run.append(rPr)
		new_run.text = anchor_text
		hyperlink.append(new_run)
		attach_to._p.append(hyperlink)

		# r = attach_to.add_run(anchor_text)
		# r._r.append(hyperlink)
		# r.font.name = "Calibri"
		# r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
		# r.font.underline = True

	else:
		add_hyperlink(paragraph=attach_to, text=anchor_text, url=target_text)



# -----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# DOCX style specific utility functions

''' return the style if exists
'''
def get_style_by_name(docx, style_name, nesting_level=0):
	for style in docx.styles:
		# style.name is the UI name (e.g., 'Normal', 'Heading 1')
		# style.element.styleId is the XML ID (e.g., 'Normal', 'Heading1')
		if style_name in (style.name, style.element.styleId):
			return style
    
	return None


''' apply a custom style to a Style/paragraph
'''
def apply_custom_style(docx, style_spec, style_name=None, paragraph=None, nesting_level=0):
	# the following elemnts are to be updated
	font, pf = None, None
	border_around = None

	# if style_name
	if style_name:
		style = get_style_by_name(docx=docx, style_name=style_name)
		if style is None:
			error(f"style [{style_name}] not found .. ")
			return

		# style exists, update with spec
		font = style.font
		pf = style.paragraph_format
		element = style._element


	elif paragraph is not None:
		font = paragraph.runs[0].font
		pf = paragraph.paragraph_format
		element = paragraph._p


	# now apply
	# ParagraphStyle
	if 'ParagraphStyle' in style_spec:
		if 'font' in style_spec['ParagraphStyle']:
			# ParagraphStyle - font
			attr_dict = style_spec['ParagraphStyle']['font']

			if font:
				for attr, value in attr_dict.items():
					# color needs special treatment
					if attr == "color":
						# value should be RGBColor or None
						font.color.rgb = value
						continue

					if hasattr(font, attr):
						setattr(font, attr, value)


			# ParagraphStyle - paragraph_format
			if 'paragraph-format' in style_spec['ParagraphStyle']:
				attr_dict = style_spec['ParagraphStyle']['paragraph-format']
				
				if pf:
					for attr, value in attr_dict.items():
						if hasattr(pf, attr):
							setattr(pf, attr, value)

	# borders
	if 'borders' in style_spec:
		if element is not None:
			set_paragraph_border(element=element, borders=style_spec['borders'])

	# backgroundcolor
	if 'backgroundcolor' in style_spec:
		if element is not None:
			set_paragraph_bgcolor(element=element, color=style_spec['backgroundcolor'])


''' process custom styles
'''
def process_custom_style(docx, style_spec, nesting_level=0):
	custom_styles = {}
	if style_spec:
		trace(f"processing custom styles from conf/style-spec.yml", nesting_level=nesting_level)
		parse_style_properties(style_spec, nesting_level=nesting_level+1)

		# iterate and apply or store
		for key, value in style_spec.items():
			if value:
				style_name = value.get('name', None)
				if style_name:
					# trace(f"customizing style [{style_name}]", nesting_level=nesting_level)
					apply_custom_style(docx=docx, style_spec=value, style_name=style_name, nesting_level=nesting_level+1)
					trace(f"customized  style [{style_name}]", nesting_level=nesting_level+1)

				else:
					# trace(f"adding custom style [{key}] to style cache", nesting_level=nesting_level)
					custom_styles[key] = value
					trace(f"added  custom style [{key}] to style cache", nesting_level=nesting_level+1)
	
	return custom_styles


''' parse style properties from yml to docx
'''
def parse_style_properties(style_spec, parent_key=None, nesting_level=0):
	if style_spec:
		for key, value in style_spec.items():
			if isinstance(value, dict):
				# If it's another dict, go deeper
				parse_style_properties(value, parent_key=key, nesting_level=nesting_level)

			else:	
				if value and value != '':
					new_value = value
					if key in DOCX_ATTR_MAP_HINT:
						# trace(f"parsing   property [{key}] with value [{value}]", nesting_level=nesting_level+1)
						new_value = map_docx_attr(key, value, parent_key=parent_key, nesting_level=nesting_level+1)
						# trace(f"parsed to property [{key}] with value [{new_value}]", nesting_level=nesting_level+1)

					elif (parent_key, key) in DOCX_ATTR_MAP_HINT:
						# trace(f"parsing   property [{parent_key, key}] with value [{value}]", nesting_level=nesting_level+1)
						new_value = map_docx_attr((parent_key, key), value, parent_key=parent_key, nesting_level=nesting_level+1)
						# trace(f"parsed to property [{parent_key, key}] with value [{new_value}]", nesting_level=nesting_level+1)

					if new_value:
						value = new_value

					style_spec[key] = value

				else:
					# style_spec.pop(key, None)
					pass

		# check for *inline-image* and process
		if 'inline-image' in style_spec:
			inline_image_list = []
			# this may be a dict for one single image or a list for multiple images
			if isinstance(style_spec['inline-image'], dict):
				inline_image_list.append(style_spec['inline-image'])

			elif isinstance(style_spec['inline-image'], list):
				for ii_dict in style_spec.get('inline-image', []):
					inline_image_list.append(ii_dict)
			
			else:
				warn(f"inline-image is neither a dict nor a list", nesting_level=nesting_level)

			style_spec['inline-image'] = []
			for ii_dict in inline_image_list:
				if 'url' in ii_dict:
					url = ii_dict.get('url')
					
					# download image
					debug(f"downloading inline image {url}", nesting_level=nesting_level+1)
					ii_image_dict = download_image(drive_service=GoogleServices().drive_api, url=url, title=None, tmp_dir=ConfigService()._temp_dir, nesting_level=nesting_level+1)

					# type background/inline
					ii_image_dict['type'] = ii_dict.get('type', 'background')

					# fit-height-to-container true/false,
					ii_image_dict['fit-height-to-container'] = ii_dict.get('fit-height-to-container', False)

					# fit-width-to-container true/false,
					ii_image_dict['fit-width-to-container'] = ii_dict.get('fit-width-to-container', False)

					# fit-height-to-container true/false,
					ii_image_dict['keep-aspect-ratio'] = ii_dict.get('keep-aspect-ratio', True)

					# extend-container-height true/false,
					ii_image_dict['extend-container-height'] = ii_dict.get('extend-container-height', False)

					# position is horizontal and vertical positions [center/left/right] [middle/top/bottom]
					ii_image_dict['position'] = ii_dict.get('position', 'center middle')

					# wrap none/parallel
					ii_image_dict['wrap'] = ii_dict.get('wrap', 'parallel')

					style_spec['inline-image'].append(ii_image_dict)
					# trace(f"downloaded  inline image {url}", nesting_level=nesting_level+1)


		# check for *page-background* and process
		if 'page-background' in style_spec:
			page_background_image_list = []
			# this may be a dict for one single image or a list for multiple images
			if isinstance(style_spec['page-background'], dict):
				page_background_image_list.append(style_spec['page-background'])

			elif isinstance(style_spec['page-background'], list):
				for pb_dict in style_spec.get('page-background', []):
					page_background_image_list.append(pb_dict)
			
			else:
				warn(f"page-background is neither a dict nor a list", nesting_level=nesting_level)

			style_spec['page-background'] = []
			for pb_dict in page_background_image_list:
				if 'url' in pb_dict:
					url = pb_dict.get('url')
					
					# download image
					debug(f"downloading inline image {url}", nesting_level=nesting_level+1)
					pb_image_dict = download_image(drive_service=GoogleServices().drive_api, url=url, title=None, tmp_dir=ConfigService()._temp_dir, nesting_level=nesting_level+1)

					# type background/inline
					pb_image_dict['type'] = 'background'

					style_spec['page-background'].append(pb_image_dict)
					# trace(f"downloaded  inline image {url}", nesting_level=nesting_level+1)



# -----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# indexes and pdf generation

''' update docx indexes by opening and closing the docx, rest is done by macros
'''
def update_indexes(docx_path, nesting_level=0):

	try:
		word = client.DispatchEx("Word.Application")
	except Exception as e:
		raise e

	try:
		docx = word.Documents.Open(docx_path)
		docx.Close()
	except Exception as e:
		raise e
	finally:
		word.Quit()


''' set docx updateFields property true
'''
def set_updatefields_true(docx_path, nesting_level=0):
	namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
	docx = Document(docx_path)
	# add child to docx.settings element
	element_updatefields = lxml.etree.SubElement(
		docx.settings.element, f"{namespace}updateFields"
	)
	element_updatefields.set(f"{namespace}val", "true")
	docx.save(docx_path)


''' given an docx file generates pdf in the given directory
'''
def generate_pdf(infile, outdir, nesting_level=0):
	# Constants for Word Export
	wdExportFormatPDF = 17
	wdExportOptimizeForPrint = 0
	wdExportAllDocument = 0
	wdExportCreateHeadingBookmarks = 1  # This enables the bookmarks

	pdf_path = infile + '.pdf'
	try:
		word = client.DispatchEx("Word.Application")
	except Exception as e:
		raise e

	try:
		docx = word.Documents.Open(infile)
		try:
			docx.ExportAsFixedFormat(pdf_path, 
				ExportFormat=wdExportFormatPDF,
				OpenAfterExport=False,
				OptimizeFor=wdExportOptimizeForPrint,
				Range=wdExportAllDocument,
				Item=0, # wdExportDocumentContent
				IncludeDocProps=True,
				KeepIRM=True,
				CreateBookmarks=wdExportCreateHeadingBookmarks # KEY PARAMETER
        )				

		except Exception as e:
			raise e

		finally:
			docx.Close()

	except Exception as e:
		raise e

	finally:
		word.Quit()


''' create table-of-contents
'''
def create_index(docx, index_type, nesting_level=0):
	paragraph = docx.add_paragraph()
	run = paragraph.add_run()

	# create a new element with attributes
	fldChar = OxmlElement('w:fldChar')
	fldChar.set(qn('w:fldCharType'), 'begin')

	instrText = OxmlElement('w:instrText')
	instrText.set(qn('xml:space'), 'preserve')

	if index_type == 'toc':
		instrText.text = 'TOC \\o "1-6" \\h \\z \\u'
	elif index_type == 'lof':
		instrText.text = 'TOC \\h \\z \\t "Figure" \\c'
	elif index_type == 'lot':
		instrText.text = 'TOC \\h \\z \\t "Table" \\c'

	fldChar2 = OxmlElement('w:fldChar')
	fldChar2.set(qn('w:fldCharType'), 'separate')
	fldChar3 = OxmlElement('w:t')
	fldChar3.text = "Right-click to update Index."
	fldChar2.append(fldChar3)

	fldChar4 = OxmlElement('w:fldChar')
	fldChar4.set(qn('w:fldCharType'), 'end')

	r_element = run._r
	r_element.append(fldChar)
	r_element.append(instrText)
	r_element.append(fldChar2)
	r_element.append(fldChar4)


''' merge another docx into this document
'''
def merge_document(placeholder, docx_path, nesting_level=0):
	# the document is in the same folder as the template, get the path
	# docx_path = os.path.abspath('{0}/{1}'.format('../conf', docx_name))
	sub_doc = Document(docx_path)

	par_parent = placeholder._p.getparent()
	index = par_parent.index(placeholder._p) + 1
	for element in sub_doc.part.element:
		element.remove_all('w:sectPr')
		par_parent.insert(index, element)
		index = index + 1



# -----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# web and drive specific utility functions

''' download an image from a web or drive url and return a dict
    {'file-path': file-path, 'file-type': file-type, 'image-height': height, 'image-width': width}
'''
def download_image(drive_service, url, title, tmp_dir, nesting_level=0):
    data = None
    if url.startswith('https://drive.google.com/'):
        data = download_file_from_drive(drive_service=drive_service, url=url, title=title, tmp_dir=tmp_dir, nesting_level=nesting_level)

    elif url.startswith('http'):
        # the file url is a normal web url
        data = download_file_from_web(url=url, tmp_dir=tmp_dir, nesting_level=nesting_level)

    else:
        warn(f"the url [{url}] is not a drive or web url", nesting_level=nesting_level)
        return None

    # if image, calculate dimensions
    if data['file-type'] in IMAGE_MIME_TYPES:
        file_path = data['file-path']
        width, height, dpi_x, dpi_y = image_meta_pillow(file_path, nesting_level=nesting_level)
        data['image-width'] = float(width / dpi_x)
        data['image-height'] = float(height / dpi_y)

    return data


''' download a file from a drive url and return a dict
    {'file-name': file-name, 'file-type': file-type, 'file-path': local_path)}
    drive file id may have many forms like
    "https://drive.google.com/open?id=1rVmH-dHciYgPwJC0EFpHIXaZ_H1j5LDu"
    "https://drive.google.com/file/d/10bjxB_yjXgtaGZWJLVCQ28a7G_PPBLk-"
'''
def download_file_from_drive(drive_service, url, title, tmp_dir, nesting_level=0):
    file_url = url.strip()

    id = file_url.replace('https://drive.google.com/', '')

    # see if it has something like 'id=' in it, then it will start after the pattern
    id = re.sub(r".*\?id=", "", id)

    # see if it has something like '/d/' in it, then it will start after the pattern
    id = re.sub(r".*/d/", "", id)

    # then it will be till end or till before the first '/'
    id = id.split('/')[0]

    # get file metadata
    file = drive_file_metadata(drive_service=drive_service, file_id=id, nesting_level=nesting_level+1)
    if not file:
        error(f"drive file [{id}] could not be accessed", nesting_level=nesting_level)
        return None
    
    if title:
        file_name = title
    else:
        file_name = file['name']
    
    file_type = file['mimeType']
    if not file_type in ALLOWED_MIME_TYPES:
        warn(f"drive url {url} is not a [{'/'.join(SUPPORTED_FILE_FORMATS)}], it is [{file_type}]", nesting_level=nesting_level)
        return None
    
    # determine file extension
    expected_extension = MIME_TYPE_TO_FILE_EXT_MAP.get(file_type, None)
    if expected_extension and not file_name.endswith(expected_extension):
        file_name = file_name + expected_extension

    local_path = f"{tmp_dir}/{file_name}"
    local_path = Path(local_path).resolve()

    # if the file is already in the local_path, we do not download it
    if Path(local_path).exists():
        trace(f"drive file existing   at: [{local_path}]", nesting_level=nesting_level)
        return {'file-name': file_name, 'file-type': file_type, 'file-path': str(local_path)}

    # finally download the file
    trace(f"downloading drive file id = [{id}]", nesting_level=nesting_level)
    try:
        download_media_from_dive(drive_service=drive_service, file_id=id, local_path=local_path, nesting_level=nesting_level+1)
        trace(f"drive file downloaded at: [{local_path}]", nesting_level=nesting_level)
        return {'file-name': file_name, 'file-type': file_type, 'file-path': str(local_path)}

    except:
        error(f"could not download : [{file_url}]", nesting_level=nesting_level)
        return None


''' download a file from a web url and return a dict
    {'file-name': file-name, 'file-type': file-type, 'file-path': local_path)}
'''
def download_file_from_web(url, tmp_dir, nesting_level=0):
    file_url = url.strip()
    file_parts = file_url.split('.')
    if len(file_parts) > 1:
        file_ext = '.' + file_parts[-1]

    if not file_ext in SUPPORTED_FILE_FORMATS:
        error(f"url {file_url} is NOT a [{'/'.join(SUPPORTED_FILE_FORMATS)}] file", nesting_level=nesting_level)
        return None

    file_name = file_url.split('/')[-1].strip()
    file_type = FILE_EXT_TO_MIME_TYPE_MAP.get(file_ext, None)

    # download pdf in url into localpath
    try:
        local_path = f"{tmp_dir}/{file_name}"
        local_path = Path(local_path).resolve()
        # if the pdf is already in the local_path, we do not download it
        if Path(local_path).exists():
            trace(f"file existing   [{file_url}]", nesting_level=nesting_level)
            # pass
        else:
            file_data = requests.get(file_url).content
            with open(local_path, 'wb') as handler:
                handler.write(file_data)

            trace(f"file downloaded [{file_url}]", nesting_level=nesting_level)

        return {'file-name': file_name, 'file-type': file_type, 'file-path': str(local_path)}
    except:
        error(f"could not download : [{file_url}]", nesting_level=nesting_level)
        return None


''' get image metadata using Pillow
'''
def image_meta_pillow(im_path, nesting_level=0):
    im = Image.open(im_path)
    width, height = im.size

    if 'dpi' in im.info:
        dpi_x, dpi_y = im.info['dpi']
    else:
        dpi_x, dpi_y = DPI, DPI

    return width, height, dpi_x, dpi_y


''' get drive file metadata dict given a file id

'''
def drive_file_metadata(drive_service, file_id, nesting_level=0):
    file = drive_service.files().get(fileId=file_id,fields="id,name,mimeType").execute()
    return file


''' download media from drive given the id to a local path
'''
def download_media_from_dive(drive_service, file_id, local_path, nesting_level=0):
    request = drive_service.files().get_media(fileId=file_id,)

    fh = io.FileIO(local_path, "wb")
    downloader = MediaIoBaseDownload(fh, request)

    done = False
    while not done:
        status, done = downloader.next_chunk()
        trace(f"Downloaded {int(status.progress() * 100)}%", nesting_level=nesting_level)

    return done



# -----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# various utility functions

''' inches to emu
'''
def emu(inches: float, nesting_level=0) -> int:
    return int(inches * EMU_PER_INCH)


''' pt to emu
'''
def pt_to_emu(pt, nesting_level=0):
    return int(pt * EMU_PER_PT)


''' given pixel size, calculate the row height in inches
	a reasonable approximation is what gsheet says 21 pixels, renders well as 12 pixel (assuming our normal text is 10-11 in size)
'''
def row_height_in_inches(pixel_size, nesting_level=0):
	return float((pixel_size) / 96)


''' get a random string
'''
def random_string(length=12, nesting_level=0):
	letters = string.ascii_uppercase
	return ''.join(random.choice(letters) for i in range(length))


''' fit width/height into a given width/height maintaining aspect ratio
'''
def fit_width_height(fit_within_width, fit_within_height, width_to_fit, height_to_fit, nesting_level=0):
	scale = min(
        fit_within_width / width_to_fit,
        fit_within_height / height_to_fit
    )

	return (
        width_to_fit * scale,
        height_to_fit * scale,
    )


''' strip delimeters for math mode
'''
def strip_math_mode_delimeters(latex_content, nesting_level=0):
	# strip SPACES
	stripped = latex_content.strip()

	# strip $
	stripped = stripped.strip('$')

	# TODO: strip \( and \)

	return stripped


''' Convert an inline picture (<wp:inline>) to a floating anchored picture (<wp:anchor>)
	positioned at page origin and behind text.
'''
def inline_to_anchored_behind(inline, width, height, nesting_level=0):
    inline_elm = inline._inline  # the <wp:inline> element

    # Rename <wp:inline> -> <wp:anchor>
    inline_elm.tag = qn("wp:anchor")

    # Anchor attributes (Word is picky about these)
    inline_elm.set("simplePos", "0")
    inline_elm.set("relativeHeight", "0")
    inline_elm.set("behindDoc", "1")
    inline_elm.set("locked", "0")
    inline_elm.set("layoutInCell", "1")
    inline_elm.set("allowOverlap", "1")
    inline_elm.set("distT", "0")
    inline_elm.set("distB", "0")
    inline_elm.set("distL", "0")
    inline_elm.set("distR", "0")

    # Build required children: simplePos, positionH, positionV
    simplePos = OxmlElement("wp:simplePos")
    simplePos.set("x", "0")
    simplePos.set("y", "0")

    positionH = OxmlElement("wp:positionH")
    positionH.set("relativeFrom", "page")
    posOffsetH = OxmlElement("wp:posOffset")
    posOffsetH.text = "0"
    positionH.append(posOffsetH)

    positionV = OxmlElement("wp:positionV")
    positionV.set("relativeFrom", "page")
    posOffsetV = OxmlElement("wp:posOffset")
    posOffsetV.text = "0"
    positionV.append(posOffsetV)

    # Insert at the beginning, BEFORE wp:extent
    # (inline originally starts with wp:extent)
    inline_elm.insert(0, simplePos)
    inline_elm.insert(1, positionH)
    inline_elm.insert(2, positionV)

    # Ensure wrapNone exists right after effectExtent
    wrapNone = OxmlElement("wp:wrapNone")

    # Find effectExtent if present; if not, create one after extent
    extent = inline_elm.find(qn("wp:extent"))
    effectExtent = inline_elm.find(qn("wp:effectExtent"))

    if effectExtent is None:
        effectExtent = OxmlElement("wp:effectExtent")
        effectExtent.set("l", "0")
        effectExtent.set("t", "0")
        effectExtent.set("r", "0")
        effectExtent.set("b", "0")
        # place it after extent
        if extent is not None:
            idx = list(inline_elm).index(extent)
            inline_elm.insert(idx + 1, effectExtent)
        else:
            inline_elm.insert(3, effectExtent)

    # Insert wrapNone immediately after effectExtent if not already there
    children = list(inline_elm)
    eff_idx = children.index(effectExtent)
    if not any(c.tag == qn("wp:wrapNone") for c in children):
        inline_elm.insert(eff_idx + 1, wrapNone)

    # IMPORTANT: set extent to full page (if you want exact A4 stretch)
    # (python-docx may have already set it based on add_picture width/height,
    # but this forces correctness)
    if extent is not None:
        extent.set("cx", str(emu(width)))
        extent.set("cy", str(emu(height)))


''' Removes/sets noChangeAspect on a:picLocks so the picture can stretch freely.
	Works for both <wp:inline> and <wp:anchor> (same underlying structure).
'''
def unlock_picture_aspect_ratio(inline_or_anchor, nesting_level=0):
    drawing = inline_or_anchor._inline.getparent()  # <w:drawing>
    # Find a:pic in the drawing
    pic = drawing.find(".//" + qn("pic:pic"))
    if pic is None:
        return

    # Non-visual picture properties: pic:nvPicPr
    nvPicPr = pic.find(qn("pic:nvPicPr"))
    if nvPicPr is None:
        return

    cNvPicPr = nvPicPr.find(qn("pic:cNvPicPr"))
    if cNvPicPr is None:
        return

    picLocks = cNvPicPr.find(qn("a:picLocks"))
    if picLocks is None:
        picLocks = OxmlElement("a:picLocks")
        cNvPicPr.append(picLocks)

    # noChangeAspect="0" means NOT locked. If it's "1" it's locked.
    picLocks.set("noChangeAspect", "0")


''' Forces the internal DrawingML transform to match width/height,
	resets scaling to 100%, and removes cropping.

	This helps when Word 'snaps back' to aspect ratio behavior or applies
	internal scaling/cropping that fights wp:extent sizing.
'''
def force_picture_transform(inline_or_anchor, width_in: float, height_in: float, nesting_level=0):
    drawing = inline_or_anchor._inline.getparent()  # <w:drawing>

    # Find picture element
    pic = drawing.find(".//" + qn("pic:pic"))
    if pic is None:
        return

    # Remove cropping if present: pic:blipFill/a:srcRect
    blipFill = pic.find(qn("pic:blipFill"))
    if blipFill is not None:
        srcRect = blipFill.find(qn("a:srcRect"))
        if srcRect is not None:
            blipFill.remove(srcRect)

    # Ensure shape properties exist: pic:spPr
    spPr = pic.find(qn("pic:spPr"))
    if spPr is None:
        spPr = OxmlElement("pic:spPr")
        pic.append(spPr)

    # Ensure transform exists: a:xfrm
    xfrm = spPr.find(qn("a:xfrm"))
    if xfrm is None:
        xfrm = OxmlElement("a:xfrm")
        # Best placed near start of spPr
        spPr.insert(0, xfrm)

    # Ensure offset exists: a:off (top-left of the shape)
    off = xfrm.find(qn("a:off"))
    if off is None:
        off = OxmlElement("a:off")
        xfrm.insert(0, off)
    off.set("x", "0")
    off.set("y", "0")

    # Ensure extent exists: a:ext (size of the shape)
    ext = xfrm.find(qn("a:ext"))
    if ext is None:
        ext = OxmlElement("a:ext")
        xfrm.append(ext)

    ext.set("cx", str(emu(width_in)))
    ext.set("cy", str(emu(height_in)))

    # Reset any scale attributes (rare, but some docs have them)
    # Some producers add chExt / chOff etc. We won’t add those; we just ensure ext is correct.


''' convert strings like '12pt' '3.00in' to Pt(12) or Inches(3.00)
'''
def str_to_size(str, what, nesting_level=0):
	allowed_units = ['pt', 'in', 'cm', 'mm']
	match = re.match(r"(\d*\.?\d+)\s*([a-zA-Z]+)", str.strip())

	# the match must return exactly two groups
	try:
		# This will raise a ValueError if there are more or fewer than 2 groups
		number, unit = match.groups()
		try:
			num = float(number)
			if unit not in allowed_units:
				warn(f"[{unit}] in [{str}] is not a valid {what} .. allowed values are [{allowed_units}]", nesting_level=nesting_level+1)
				return None
			
			if unit == 'pt':
				return Pt(num)
			
			if unit == 'in':
				return Inches(num)
			
			if unit == 'cm':
				return Cm(num)

		except:
			warn(f"[{str}] is not a valid {what} .. first part must be a number", nesting_level=nesting_level+1)
			return None

	except ValueError:
		warn(f"[{str}] is not a valid {what}", nesting_level=nesting_level+1)
		return None

	return Pt(float(num))


''' convert strings like '1pt double #222222 0.25in' to border {"sz": , "val": , "color": , "space": , "shadow": }
	size is in 1/8 pt units
	val is any of dotted/dashed/single/thick/triple/double/none
	space is space/padding between text and border in pt
'''
def str_to_border(str, what, nesting_level=0):
	try:
		# we are expecting at least four parts (sz, val, color, space)
		sz, val, color, space = str.split()

	except ValueError:
		warn(f"[{str}] is not a valid {what} ... expecting four parts sz val color space", nesting_level=nesting_level+1)
		return None

	# sz
	sz = str_to_size(str=sz, what=what, nesting_level=nesting_level)
	if sz is not None:
		sz = sz.pt * 8
	else:
		return None

	# val
	allowed_vals = ['dotted', 'dashed', 'single', 'thick', 'triple', 'double', 'none']
	if val not in allowed_vals:
		warn(f"[{val}] is not valid in {what} ... valid values are {allowed_vals}", nesting_level=nesting_level+1)
		return None

	# color
	color = color.lstrip("#")
	if not is_valid_hex(color, 6):
		warn(f"[{color}] is not valid in {what} ... not a valid 6 digit hex color", nesting_level=nesting_level+1)
		return None

	# space
	space = str_to_size(str=space, what=what, nesting_level=nesting_level)
	if space:
		space = space.pt
	else:
		return None
	
	return {"sz": int(sz), "val": val, "color": color, "space": int(space), "shadow": None}


''' conver rgb colors like '#RRGGBB' to RGBColor
'''
def rgb_from_hex(str, what, nesting_level=0):
    hexstr = str.lstrip("#")
    if is_valid_hex(hexstr, 6):
        return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16),)
    else:
    	warn(f"[{str}] is not a valid {what} ... not a valid 6 digit hex color", nesting_level=nesting_level+1)


''' is the string a valid hex-string
'''
def is_valid_hex(str, digits, nesting_level=0):
	pattern = rf'^[0-9a-fA-F]{{{digits}}}$'
	return bool(re.fullmatch(pattern, str))


''' pretty print element xml
'''
def print_xml(element, nesting_level=0):
	# Convert your element to a string first (using ET or lxml)
	xml_str = etree.tostring(element, encoding='unicode', pretty_print=True)


''' map attribute key/value to a modified ke/value
'''
def map_docx_attr(attr_key, attr_value, parent_key, nesting_level=0):
	if attr_key in DOCX_ATTR_MAP_HINT:
		obj = DOCX_ATTR_MAP_HINT[attr_key]
		mapper = obj['lambda']

		# if the mapper is a dict
		if isinstance(mapper, dict):
			# trace(f"[{attr_key}] mapper is a dict", nesting_level=nesting_level)
			if attr_value in mapper:
				return mapper[attr_value]
			
			else:
				warn(f"[{attr_value}] in [{attr_key}] is not .. allowed values are {list(mapper.keys())}", nesting_level=nesting_level)

		# of the mapper is a function
		if isinstance(mapper, types.FunctionType):
			# trace(f"[{attr_key}] mapper is a function [{mapper}]", nesting_level=nesting_level)
			return mapper(attr_value, what=attr_key, nesting_level=nesting_level)
		
	return attr_value
			


# -----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# various utility data

# mml2omml_stylesheet_path = '../conf/MML2OMML_15.XSL'
mml2omml_stylesheet_path = '../conf/MML2OMML_16.XSL'

# emu per inch
EMU_PER_INCH = 914500

# emu per pt
EMU_PER_PT = 12700

# default DPI
DPI = 72

# height offset for full page image extracted from pdf
PDF_PAGE_HEIGHT_OFFSET = 0.5

# 0-based gsheet column number to column letter map
COLUMNS = [ 'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z',
			'AA', 'AB', 'AC', 'AD', 'AE', 'AF', 'AG', 'AH', 'AI', 'AJ', 'AK', 'AL', 'AM', 'AN', 'AO', 'AP', 'AQ', 'AR', 'AS', 'AT', 'AU', 'AV', 'AW', 'AX', 'AY', 'AZ',
			'BA', 'BB', 'BC', 'BD', 'BE', 'BF', 'BG', 'BH', 'BI', 'BJ', 'BK', 'BL', 'BM', 'BN', 'BO', 'BP', 'BQ', 'BR', 'BS', 'BT', 'BU', 'BV', 'BW', 'BX', 'BY', 'BZ']



# -----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# gsheet to odt constants and type mapping

HEADING_TO_LEVEL = {
	'Heading 1': {'outline-level': 1},
	'Heading 2': {'outline-level': 2},
	'Heading 3': {'outline-level': 3},
	'Heading 4': {'outline-level': 4},
	'Heading 5': {'outline-level': 5},
	'Heading 6': {'outline-level': 6},
	'Heading 7': {'outline-level': 7},
	'Heading 8': {'outline-level': 8},
	'Heading 9': {'outline-level': 9},
	'Heading 10': {'outline-level': 10},
}

LEVEL_TO_HEADING = [
	'Title',
	'Heading 1',
	'Heading 2',
	'Heading 3',
	'Heading 4',
	'Heading 5',
	'Heading 6',
	'Heading 7',
	'Heading 8',
	'Heading 9',
	'Heading 10',
]

CELL_VALIGN_MAP = {
	'TOP': WD_CELL_VERTICAL_ALIGNMENT.TOP,
	'MIDDLE': WD_CELL_VERTICAL_ALIGNMENT.CENTER,
	'BOTTOM': WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
}

TEXT_HALIGN_MAP = {
	'LEFT': WD_PARAGRAPH_ALIGNMENT.LEFT,
	'CENTER': WD_PARAGRAPH_ALIGNMENT.CENTER,
	'RIGHT': WD_PARAGRAPH_ALIGNMENT.RIGHT,
	'JUSTIFY': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
}

WRAP_STRATEGY_MAP = {
	'OVERFLOW': 'no-wrap', 'CLIP': 'no-wrap', 'WRAP': 'wrap'
}

GSHEET_OXML_BORDER_MAPPING = {
	'DOTTED': 'dotted',
	'DASHED': 'dashed',
	'SOLID': 'single',
	'SOLID_MEDIUM': 'thick',
	'SOLID_THICK': 'triple',
	'DOUBLE': 'double',
	'NONE': 'none'
}



# -----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# allowed file and mimetype related

SUPPORTED_FILE_FORMATS = ['.pdf', '.png', '.jpg', '.gif', '.webp']

ALLOWED_MIME_TYPES = ['application/pdf', 'image/png', 'image/jpeg', 'image/gif', 'image/webp']

IMAGE_MIME_TYPES = ['image/png', 'image/jpeg', 'image/gif', 'image/webp']

MIME_TYPE_TO_FILE_EXT_MAP = {
    'application/pdf': '.pdf', 
    'image/png': '.png', 
    'image/jpeg': '.jpg', 
    'image/gif': '.gif', 
    'image/webp': '.webp'
}

FILE_EXT_TO_MIME_TYPE_MAP = {
    '.pdf': 'application/pdf', 
    '.png': 'image/png', 
    '.jpg': 'image/jpeg', 
    '.gif': 'image/gif', 
    '.webp': 'image/webp' 
}



# -----------------------------------------------------------------------------------------------------------------------------------------------------------------------------
# custom style map specific

DOCX_ATTR_MAP_HINT = {
	'color':				{'lambda': rgb_from_hex},
	'highlight_color':		{'lambda': rgb_from_hex},
	# 'backgroundcolor':	{'lambda': rgb_from_hex},
	'alignment': 			{'lambda': TEXT_HALIGN_MAP},
	'verticalalign':		{'lambda': CELL_VALIGN_MAP},
	'size': 				{'lambda': str_to_size},
	'left_indent':      	{'lambda': str_to_size},
	'right_indent':     	{'lambda': str_to_size},
	'space_before':     	{'lambda': str_to_size},
	'space_after':    		{'lambda': str_to_size},
	('borders', 'start'):   {'lambda': str_to_border},
	('borders', 'top'):     {'lambda': str_to_border},
	('borders', 'end'):     {'lambda': str_to_border},
	('borders', 'bottom'):  {'lambda': str_to_border},
	('padding', 'start'):   {'lambda': str_to_size},
	('padding', 'top'):     {'lambda': str_to_size},
	('padding', 'end'):     {'lambda': str_to_size},
	('padding', 'bottom'):  {'lambda': str_to_size},
}

STYLE_PROPERTY_MAP = {
	"ParagraphStyle": 
	{
		"font":
		[
			"name",
			"size",
			"color",
			"bold",
			"italic",
			"underline",
			"strike",
			"double_strike",
			"highlight_color",
			"all_caps",
			"small_caps",
			"subscript",
			"superscript",
			"complex_script",
			"cs_bold",
			"cs_italic",
			"emboss",
			"imprint",
			"outline",
			"shadow",
		],
		"paragraph_format":
		[
			# A member of the WD_PARAGRAPH_ALIGNMENT enumeration specifying the justification setting for this paragraph.
			"alignment",
			# Length value specifying the relative difference in indentation for the first line of the paragraph.
			"first_line_indent",
			# True if the paragraph should be kept “in one piece” and not broken across a page boundary when the document is rendered.
			"keep_together",
			"keep_with_next",
			# Length value specifying the space between the left margin and the left side of the paragraph.
			"left_indent",
			# float or Length value specifying the space between baselines in successive lines of the paragraph.
			"line_spacing",
			# A member of the WD_LINE_SPACING enumeration indicating how the value of line_spacing should be interpreted.
			"line_spacing_rule",
			"page_break_before",
			"right_indent",
			# Length value specifying the spacing to appear between this paragraph and the subsequent paragraph.
			"space_after",
			# Length value specifying the spacing to appear between this paragraph and the prior paragraph.
			"space_before",
			# TabStops object providing access to the tab stops defined for this paragraph format
			"tab_stops",
			# True if the first and last lines in the paragraph remain on the same page as the rest of the paragraph when Word repaginates the document.
			"widow_control",
		],
	},
}
