#!/usr/bin/env python

from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

A4_W = Inches(8.27)
A4_H = Inches(11.69)
A4_W_IN = 8.27
A4_H_IN = 11.69

def set_a4_no_margins(section):
    section.page_width = A4_W
    section.page_height = A4_H
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)


EMU_PER_INCH = 914400  # Word uses EMU units

def _emu(inches: float) -> int:
    return int(inches * EMU_PER_INCH)

def make_inline_picture_anchored_behind(inline, page_width_in=8.27, page_height_in=11.69):
    """
    Convert a python-docx inline picture into a floating anchored picture,
    placed at the top-left of the page and behind text.

    This edits the existing <wp:inline> so we preserve the valid picture structure.
    """
    inline_elm = inline._inline  # the <wp:inline> element

    # Rename <wp:inline> -> <wp:anchor>
    inline_elm.tag = qn("wp:anchor")

    # Anchor attributes (Word is picky about these)
    inline_elm.set("simplePos", "0")
    inline_elm.set("relativeHeight", "0")
    inline_elm.set("behindDoc", "1")
    inline_elm.set("locked", "0")
    inline_elm.set("layoutInCell", "1")
    inline_elm.set("allowOverlap", "1")
    inline_elm.set("distT", "0")
    inline_elm.set("distB", "0")
    inline_elm.set("distL", "0")
    inline_elm.set("distR", "0")

    # Build required children: simplePos, positionH, positionV
    simplePos = OxmlElement("wp:simplePos")
    simplePos.set("x", "0")
    simplePos.set("y", "0")

    positionH = OxmlElement("wp:positionH")
    positionH.set("relativeFrom", "page")
    posOffsetH = OxmlElement("wp:posOffset")
    posOffsetH.text = "0"
    positionH.append(posOffsetH)

    positionV = OxmlElement("wp:positionV")
    positionV.set("relativeFrom", "page")
    posOffsetV = OxmlElement("wp:posOffset")
    posOffsetV.text = "0"
    positionV.append(posOffsetV)

    # Insert at the beginning, BEFORE wp:extent
    # (inline originally starts with wp:extent)
    inline_elm.insert(0, simplePos)
    inline_elm.insert(1, positionH)
    inline_elm.insert(2, positionV)

    # Ensure wrapNone exists right after effectExtent
    wrapNone = OxmlElement("wp:wrapNone")

    # Find effectExtent if present; if not, create one after extent
    extent = inline_elm.find(qn("wp:extent"))
    effectExtent = inline_elm.find(qn("wp:effectExtent"))

    if effectExtent is None:
        effectExtent = OxmlElement("wp:effectExtent")
        effectExtent.set("l", "0")
        effectExtent.set("t", "0")
        effectExtent.set("r", "0")
        effectExtent.set("b", "0")
        # place it after extent
        if extent is not None:
            idx = list(inline_elm).index(extent)
            inline_elm.insert(idx + 1, effectExtent)
        else:
            inline_elm.insert(3, effectExtent)

    # Insert wrapNone immediately after effectExtent if not already there
    children = list(inline_elm)
    eff_idx = children.index(effectExtent)
    if not any(c.tag == qn("wp:wrapNone") for c in children):
        inline_elm.insert(eff_idx + 1, wrapNone)

    # IMPORTANT: set extent to full page (if you want exact A4 stretch)
    # (python-docx may have already set it based on add_picture width/height,
    # but this forces correctness)
    if extent is not None:
        extent.set("cx", str(_emu(page_width_in)))
        extent.set("cy", str(_emu(page_height_in)))


def add_background_image_to_header(header, image_path):
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    run = p.add_run()

    # Insert as inline (valid)
    inline = run.add_picture(image_path, width=Inches(A4_W_IN), height=Inches(A4_H_IN))

    # Convert to anchored behind text (still valid if done carefully)
    make_inline_picture_anchored_behind(inline, page_width_in=A4_W_IN, page_height_in=A4_H_IN)


def clear_header_footer(hf):
    # Remove all paragraphs in header/footer so we can set our own cleanly
    for p in list(hf.paragraphs):
        p._element.getparent().remove(p._element)

def build_first_page_header_footer(section):
    """
    Customize page 1 header/footer here.
    """
    # FIRST PAGE header/footer objects
    h = section.first_page_header
    f = section.first_page_footer

    clear_header_footer(h)
    clear_header_footer(f)

    # Example content (replace with yours)
    h.add_paragraph("FIRST PAGE HEADER")
    f.add_paragraph("FIRST PAGE FOOTER")

def build_odd_header_footer(section):
    """
    Customize ODD (primary) header/footer template here.
    """
    h = section.header
    f = section.footer

    clear_header_footer(h)
    clear_header_footer(f)

    h.add_paragraph("ODD PAGE HEADER")
    f.add_paragraph("ODD PAGE FOOTER")

def build_even_header_footer(section):
    """
    Customize EVEN header/footer template here.
    """
    h = section.even_page_header
    f = section.even_page_footer

    clear_header_footer(h)
    clear_header_footer(f)

    h.add_paragraph("EVEN PAGE HEADER")
    f.add_paragraph("EVEN PAGE FOOTER")

def create_doc_with_5_background_pages(
    image_paths,  # list of 5 A4 images
    output_path="output.docx",
):
    assert len(image_paths) == 5, "Need exactly 5 image paths."

    doc = Document()

    # Global setting: enable different odd/even headers across the document
    doc.settings.odd_and_even_pages_header_footer = True

    # -------- Section 1 (Page 1) --------
    s1 = doc.sections[0]
    set_a4_no_margins(s1)

    # Enable "Different first page" for the first section so page 1 can use first_page_header/footer
    s1.different_first_page_header_footer = True

    # Build content for first page header/footer
    build_first_page_header_footer(s1)

    # Put background image for page 1 in first-page header
    add_background_image_to_header(s1.first_page_header, image_paths[0])

    # Ensure page 1 has at least one body paragraph so the page exists
    doc.add_paragraph("")

    # -------- Sections 2..5 (Pages 2..5) --------
    # These pages should use odd/even header/footer. We'll build both for each section,
    # then inject the background into the header that matches the page parity.
    for i in range(2, 6):  # page numbers 2..5
        sec = doc.add_section(WD_SECTION.NEW_PAGE)
        set_a4_no_margins(sec)

        # For these sections, we do NOT need different_first_page_header_footer
        sec.different_first_page_header_footer = False

        # Make sure each section does NOT link headers/footers to previous,
        # otherwise the background images would get inherited.
        sec.header.is_linked_to_previous = False
        sec.footer.is_linked_to_previous = False
        sec.even_page_header.is_linked_to_previous = False
        sec.even_page_footer.is_linked_to_previous = False

        # Build odd/even templates
        build_odd_header_footer(sec)
        build_even_header_footer(sec)

        # Add the background image to the header Word will use for that page
        img = image_paths[i - 1]  # image index 1..4
        if i % 2 == 0:
            add_background_image_to_header(sec.even_page_header, img)
        else:
            add_background_image_to_header(sec.header, img)

        # body content so the page exists
        doc.add_paragraph("")

    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    images = [
        "../../out/tmp/a4-page-bg-01.jpg",
        "../../out/tmp/3.__JV Agreement-a__0001-1.jpg",
        "../../out/tmp/3.__JV Agreement-a__0001-2.jpg",
        "../../out/tmp/3.__JV Agreement-a__0001-3.jpg",
        "../../out/tmp/3.__JV Agreement-a__0001-4.jpg",
    ]
    create_doc_with_5_background_pages(images, "five_pages_backgrounds.docx")
