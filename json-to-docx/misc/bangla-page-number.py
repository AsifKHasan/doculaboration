import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_bangla_page_number_header():
    # 1. Initialize Document
    doc = docx.Document()
    
    # 2. Access the header of the default first section
    section = doc.sections[0]
    header = section.header
    
    # 3. Create a paragraph in the header and align it
    # We use an empty string so we can manually construct custom runs
    p = header.paragraphs[0]
    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
    
    # 4. Add static prefix text "পৃষ্ঠা নং: " inside a Bengali-configured run
    prefix_run = p.add_run("পৃষ্ঠা নং: ")
    set_run_to_bangla(prefix_run)
    
    # 5. Inject the Dynamic Page Field utilizing the Bengali Locale (LCID: 1093)
    # This acts identically to creating `{ PAGE \l 1093 }` inside Word
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), r'PAGE \l 1093')
    
    # Create the internal text visual wrapper inside the field code
    r = OxmlElement('w:r')
    set_run_to_bangla(r, is_raw_xml=True) # Ensure the field output inherits Bengali rules
    
    t = OxmlElement('w:t')
    t.text = "১"  # Placeholder value; Word dynamically updates this on launch
    
    r.append(t)
    fldSimple.append(r)
    p._p.append(fldSimple)
    
    # 6. Add some core body text to confirm layout structure
    doc.add_paragraph("এটি একটি নমুনা ডকুমেন্ট। (Sample core document text written in Bengali.)")
    
    # Save output
    doc.save("bangla_dynamic_page.docx")
    print("Document created successfully with Bengali dynamic page numbering.")

def set_run_to_bangla(run_obj, is_raw_xml=False):
    """
    Injects Complex Text Layout flags directly into MS Word xml structure 
    to force rendering engines to evaluate numeral formats via South Asian localization.
    """
    # If it's a python-docx Run object, get its underlying XML element; otherwise use it directly
    r_element = run_obj._r if not is_raw_xml else run_obj
    
    # Access or create Run Properties (<w:rPr>)
    rPr = r_element.get_or_add_rPr()
    
    # Define Complex Text Layout fonts and system lang identifiers
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'en-US')         # Standard Latin Fallback
    lang.set(qn('w:bidi'), 'bn-BD')        # Bi-directional / South Asian script flag
    lang.set(qn('w:bidi'), 'bn-BD') 
    
    # Set explicit font formatting targets for Complex layouts
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:csa'), 'Kalpurush')   # Standard Bengali Unicode Font
    rFonts.set(qn('w:bidi'), 'Kalpurush')  # Apply font constraints to layout boundaries
    
    rPr.append(lang)
    rPr.append(rFonts)

if __name__ == "__main__":
    create_bangla_page_number_header()