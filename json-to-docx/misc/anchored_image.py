#!/usr/bin/env python

import random

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches

EMU_PER_PT = 12700

''' pt to emu
'''
def _pt_to_emu(pt):
    return int(pt * EMU_PER_PT)


''' insert anchored image in a cell with margins
'''
def insert_cell_image(cell, image_path, width, position='center bottom', margin_pt=2, offsets_pt=(2, 2), nesting_level=0):
    relative_h = 'column'
    relative_v = 'paragraph'

    # Optional offsets (in EMU)
    offX = _pt_to_emu(offsets_pt[0])
    offY = _pt_to_emu(offsets_pt[1])

    # If margin is a single number, apply to all sides. 
    if isinstance(margin_pt, int):
        m = {'t': _pt_to_emu(margin_pt), 'b': _pt_to_emu(margin_pt), 'l': _pt_to_emu(margin_pt), 'r': _pt_to_emu(margin_pt)}

    # If it's a dict, use specific values.
    else:
        m = {'t': _pt_to_emu(margin_pt.get('t', 0)), 'b': _pt_to_emu(margin_pt.get('b', 0)), 
             'l': _pt_to_emu(margin_pt.get('l', 0)), 'r': _pt_to_emu(margin_pt.get('r', 0))}

    mapping = {
        'left top':      {'h': 'left',   'v': 'top'},
        'center top':    {'h': 'center', 'v': 'top'},
        'right top':     {'h': 'right',  'v': 'top'},
        'left middle':   {'h': 'left',   'v': 'center'},
        'center middle': {'h': 'center', 'v': 'center'},
        'right middle':  {'h': 'right',  'v': 'center'},
        'left bottom':   {'h': 'left',   'v': 'bottom'},
        'center bottom': {'h': 'center', 'v': 'bottom'},
        'right bottom':  {'h': 'right',  'v': 'bottom'},
    }
    
    aligns = mapping.get(position, mapping['center middle'])
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    shape = run.add_picture(image_path, width=Inches(width))
    inline = shape._inline
    rId = inline.graphic.graphicData.pic.blipFill.blip.embed
    cx, cy = inline.extent.cx, inline.extent.cy
    pic_id = random.randint(1000, 10000)
    name = f"Picture-{pic_id}"

    # Logic: distT, distB, distL, distR control the "text wrap buffer"
    anchor_xml = f'''
	<wp:anchor distT="{m['t']}" distB="{m['b']}" distL="{m['l']}" distR="{m['r']}"
            simplePos="0" relativeHeight="251658240" behindDoc="0"
            locked="0" layoutInCell="1" allowOverlap="1"
            {nsdecls('wp', 'a', 'pic', 'r')}>
		<wp:simplePos x="0" y="0" />
        <wp:positionH relativeFrom="{relative_h}">
            <wp:align>{aligns['h']}</wp:align>
        </wp:positionH>
        <wp:positionV relativeFrom="{relative_v}">
            <wp:align>{aligns['v']}</wp:align>
        </wp:positionV>
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
		<wp:wrapSquare wrapText="bothSides" />
		<wp:docPr id="{pic_id}" name="{name}" />
		<wp:cNvGraphicFramePr>
			<a:graphicFrameLocks
				xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
				noChangeAspect="1" />
		</wp:cNvGraphicFramePr>
		<a:graphic
			xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
			<a:graphicData
				uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
				<pic:pic
					xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
					<pic:nvPicPr>
						<pic:cNvPr id="{pic_id}" name="{name}" />
						<pic:cNvPicPr />
					</pic:nvPicPr>
					<pic:blipFill>
						<a:blip r:embed="{rId}">
						</a:blip>
						<a:stretch>
							<a:fillRect />
						</a:stretch>
					</pic:blipFill>
					<pic:spPr>
						<a:xfrm>
							<a:off x="0" y="0" />
							<a:ext cx="{cx}" cy="{cy}" />
						</a:xfrm>
						<a:prstGeom prst="rect">
							<a:avLst />
						</a:prstGeom>
					</pic:spPr>
				</pic:pic>
			</a:graphicData>
		</a:graphic>
	</wp:anchor>
    '''
    
    anchor = parse_xml(anchor_xml)
    inline.getparent().replace(inline, anchor)


doc = Document()
table = doc.add_table(rows=1, cols=1)
cell = table.cell(0, 0)
cell.text = "Text wrapping with a custom margin buffer. " * 50

insert_cell_image(
    cell, 
    '../../out/tmp/flag--bd-4x3-32px.png', 
    width=1, 
    position='center middle', 
    offsets_pt=(2, 2),
    margin_pt=2
)

doc.save('../../out/margin_control.docx')