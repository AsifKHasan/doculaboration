from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_document_language(document, lang="bn-BD"):
    """
    Set default document language.
    """
    styles = document.styles

    for style in styles:
        try:
            rPr = style.element.get_or_add_rPr()
            lang_el = rPr.find(qn('w:lang'))

            if lang_el is None:
                lang_el = OxmlElement('w:lang')
                rPr.append(lang_el)

            lang_el.set(qn('w:val'), lang)
            lang_el.set(qn('w:eastAsia'), lang)
            lang_el.set(qn('w:bidi'), lang)

        except Exception:
            pass


def add_page_number(run):
    """
    Add a dynamic PAGE field.
    """

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' PAGE '

    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')

    # Placeholder text shown before field update
    text = OxmlElement('w:t')
    text.text = "1"

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    r = run._r
    r.append(fld_begin)
    r.append(instr_text)
    r.append(fld_separate)
    r.append(text)
    r.append(fld_end)


# ------------------------------------------------------------------
# Create document
# ------------------------------------------------------------------

doc = Document()

# Set document language to Bengali
set_document_language(doc, "bn-BD")

# Header
section = doc.sections[0]
header = section.header

p = header.paragraphs[0]

run = p.add_run("পৃষ্ঠা ")
add_page_number(run)

# Add content so document has multiple pages
for i in range(500):
    doc.add_paragraph(
        f"এটি পরীক্ষামূলক অনুচ্ছেদ নম্বর {i+1}"
    )

doc.save("bangla_page_number.docx")